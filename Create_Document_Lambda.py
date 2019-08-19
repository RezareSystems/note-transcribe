""" Produce Word Document transcriptions using the automatic speech recognition from AWS Transcribe. """

#import sys
#!{sys.executable} -m pip install python-docx

from docx import Document
from docx.shared import Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, datetime
import boto3
import tempfile

s3client = boto3.client('s3')
s3resource = boto3.resource('s3')

def convert_time_stamp(n):
    """ Function to help convert timestamps from s to H:M:S """
    ts = datetime.timedelta(seconds=float(n))
    ts = ts - datetime.timedelta(microseconds=ts.microseconds)
    return str(ts)


def write(bucket, file, job_name, **kwargs):
    """ Write a transcript from the .json transcription file. """

    # Initiate Document
    document = Document()
    # A4 Size
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)
    # Font
    font = document.styles['Normal'].font
    font.name = 'Calibri'

    # Load Transcription output
    #data = json.load(open(file, 'r', encoding='utf-8'))

    content_object = s3resource.Object(bucket, file)
    file_content = content_object.get()['Body'].read().decode('utf-8')
    data = json.loads(file_content)

    # Document title and intro
    #title = f"Transcription of {data['jobName']}"
    title = f"Meeting Transcription"
    document.add_heading(title, level=1)
    
    # Meeting name
    meetingName = f"Weekly catch-up ScribeMate project"
    paragragh = document.add_paragraph()
    paragragh.add_run("Meeting Name:  ").bold = True
    paragragh.add_run(meetingName)
    
    # Date && time
    date = datetime.datetime.now()
    paragragh = document.add_paragraph()
    paragragh.add_run("Date:  ").bold = True
    paragragh.add_run(date.strftime('%A %d %B %Y')) 
    
    paragragh = document.add_paragraph()
    paragragh.add_run("Time:  ").bold = True
    paragragh.add_run(date.strftime('%X')) 
                 
    
    # find list of all speaker labels, and give default names
    speakerDict = {}

    j = 0
    for segment in data['results']['speaker_labels']['segments']:
        spk = segment['speaker_label']
        if not spk in speakerDict:
             j += 1
             speakerDict[spk] = 'Speaker ' + str(j)
     
    # substitute actual names for speakers if we find "my name is ..." in the content
    k = 0            
    for item in data['results']['items']:
        if item['type'] == 'pronunciation': 
            if len(item['alternatives']) > 0:
                current_word = dict()
                confidence_scores = []
                for score in item['alternatives']:
                    confidence_scores.append(score['confidence'])
                    for alternative in item['alternatives']:
                        if alternative['confidence'] == max(confidence_scores):
                            current_word = alternative.copy()
                theWord = current_word['content']            
                if k == 3:
                    for segment in data['results']['speaker_labels']['segments']:
                        for item1 in segment['items']:
                            if item1['start_time'] == item['start_time']:
                                speakerDict[segment['speaker_label']] = theWord
                    k = 0                              
                if theWord == 'my' or theWord == 'My':
                    k = 1           
                if theWord == 'name' and k == 1: 
                    k = 2            
                if theWord == 'is' and k == 2:  
                    k = 3                                              

     
    # list of attendees 
    paragragh = document.add_paragraph()
    paragragh.add_run("Attendees:  ").bold = True  
    
    k = 0
    for item in speakerDict.values():
        if k > 0:
            paragragh.add_run(", ")
        paragragh.add_run(item) 
        k+=1      
      
        
    # Set thresholds for formatting later
    threshold_for_grey = 0.98
       
    # Add paragraph for spacing
    document.add_paragraph()

    # Intro
    document.add_paragraph().add_run('Notes').bold = True
    #document.add_paragraph(f"Grey text has less than {int(threshold_for_grey * 100)}% confidence.")

    # Process and display transcript by speaker segments
    lastSpeaker = ''
    for segment in data['results']['speaker_labels']['segments']:
        # If there is content in the segment
        if len(segment['items']) > 0:
            # Add a row, write the time and speaker  
            spk = speakerDict[str(segment['speaker_label'])]
            if spk != lastSpeaker:
                #paragragh = document.add_paragraph(convert_time_stamp(segment['start_time']) + "  " + str(segment['speaker_label']))
                document.add_paragraph().add_run(spk + ":").bold = True
                paragragh = document.add_paragraph()
                
            lastSpeaker = spk
            # Segments group individual word results by speaker. They are cross-referenced by time.
            # For each word in the segment...
            for word in segment['items']:
                # Run through the word results and get the corresponding result
                for result in data['results']['items']:
                    if result['type'] == 'pronunciation':
                        if result['start_time'] == word['start_time']:

                            # Get the word with the highest confidence
                            if len(result['alternatives']) > 0:
                                current_word = dict()
                                confidence_scores = []
                                for score in result['alternatives']:
                                    confidence_scores.append(score['confidence'])
                                for alternative in result['alternatives']:
                                    if alternative['confidence'] == max(confidence_scores):
                                        current_word = alternative.copy()

                                # Write and format the word
                                run = paragragh.add_run(' ' + current_word['content'])
                                if float(current_word['confidence']) < threshold_for_grey:
                                    font = run.font
                                    font.color.rgb = RGBColor(150, 150, 150)

                                # If the next item is punctuation, add it
                                try:
                                    if data['results']['items'][data['results']['items'].index(result) + 1]['type'] == 'punctuation':                            
                                        run = paragragh.add_run(data['results']['items'][data['results']['items'].index(result) + 1]['alternatives'][0]['content'])
                                # Occasional IndexErrors encountered
                                except:
                                    pass

    # Save
    filename = kwargs.get('save_as', f"{data['jobName']}.docx")
    temp = tempfile.TemporaryFile()
    document.save(temp)
    temp.seek(0)
    s3client.put_object(Body = temp.read(), Bucket = bucket, Key=job_name + '.docx')
    print(f"Transcript {filename} writen.")
    
def lambda_handler(event, context):
    i = 0
    for record in event['Records']:
        # transcribe-processed-bucket
        bucket = record['s3']['bucket']['name']
        # find file name
        key = record['s3']['object']['key']
        extension = key.split('.')[-1]
        job_name = key.split('.')[0]
        i += 1
        
        if extension == 'json':
            docxFileSaveAs = 'https://s3.amazonaws.com/' + bucket + '/' +job_name + ".docx"
            write(bucket, key, job_name, save_as=docxFileSaveAs)        
    return {
        'statusCode': 200,
        'body': json.dumps(str(i) + ' Files Processed')
}
